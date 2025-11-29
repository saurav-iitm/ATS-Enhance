from flask import Flask, render_template, request, jsonify, send_file
from werkzeug.utils import secure_filename
import google.generativeai as genai
import os, re, tempfile
from dotenv import load_dotenv
from io import BytesIO
from docx import Document
from reportlab.pdfgen import canvas
import PyPDF2

# ===================================================== #
#                   LOAD GEMINI KEY                     #
# ===================================================== #
load_dotenv()

genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
if not os.getenv("GEMINI_API_KEY"):
    raise Exception("❌ ERROR: Please set GEMINI_API_KEY inside .env")

app = Flask(__name__)

# ===================================================== #
#                  ALLOWED FORMATS                      #
# ===================================================== #
ALLOWED_EXTENSIONS = {"pdf", "docx", "txt", "rtf", "html"}

def allowed_file(name: str) -> bool:
    return "." in name and name.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

# ===================================================== #
#                KEYWORD EXTRACTOR                      #
# ===================================================== #
STOPWORDS = {
    "the","and","or","a","an","in","on","for","to","of","with",
    "is","are","this","that","by","as","be","at","from","it",
    "was","were","will","can"
}

def extract_keywords(txt: str) -> set:
    words = re.findall(r"[a-zA-Z]+", txt.lower())
    return {w for w in words if len(w) > 3 and w not in STOPWORDS}

# ===================================================== #
#                    ATS SCORE                          #
# ===================================================== #
def compute_ats_score(resume_text: str, job_text: str):
    job_kw = extract_keywords(job_text)
    res_kw = extract_keywords(resume_text)

    matched = job_kw & res_kw
    missing = job_kw - res_kw

    score = int((len(matched) / max(1, len(job_kw))) * 100)

    return {
        "score": score,
        "matched": list(matched),
        "missing": list(missing),
    }

# ===================================================== #
#          AI RESUME REWRITER (STRUCTURE-AWARE TEXT)    #
# ===================================================== #
def ai_rewrite_resume(original_text, job_text):

    prompt = f"""
Rewrite the following resume into a clean, professional, ATS-friendly format.

MANDATORY RULES:
- DO NOT use stars (*), bullets like “•”, markdown, or HTML tags.
- DO NOT output bold, italics, or special symbols.
- KEEP a clean plain-text format.
- Only headings should exist, such as:
  Summary
  Skills
  Experience
  Projects
  Education
  Achievements
  Leadership
- Improve writing but do NOT invent fake experience.
- Insert job description keywords naturally.
- Maintain similar length to the original.

FORMAT TEMPLATE TO FOLLOW EXACTLY:

[Full Name]
[Location]
[Phone] | [Email]

Summary
[2–4 lines]

Skills
[Use clean hyphens or commas only]

Experience
[Job Title] | [Company] | [Dates] | [Location]
- Clean plain text bullet sentence
- Another responsibility

Projects
[Project Title]
- One line description

Education
[Degree] | [University] | [Years]

Achievements
- Achievement 1
- Achievement 2

Leadership
- Leadership role description

Rewrite the resume using only this clean professional structure.

RESUME:
{original_text}

JOB DESCRIPTION:
{job_text}
"""

    model = genai.GenerativeModel("gemini-2.0-flash")
    out = model.generate_content(prompt)
    return out.text.strip()


# ===================================================== #
#               KEYWORD HIGHLIGHTER (HTML)              #
# ===================================================== #
def highlight_keywords(text, job_text):
    job_kw = extract_keywords(job_text)

    for word in sorted(job_kw, key=len, reverse=True):
        text = re.sub(
            fr"(?i)\b{word}\b",
            rf"<span class='added'>{word}</span>",
            text
        )
    return text


# ===================================================== #
#               SAFE FILE TEXT EXTRACTION               #
# ===================================================== #
def extract_pdf(f) -> str:
    f.stream.seek(0)
    reader = PyPDF2.PdfReader(f)
    txt = ""
    for p in reader.pages:
        try:
            txt += p.extract_text() + "\n"
        except:
            pass
    return txt.strip()

def extract_docx(f) -> str:
    f.stream.seek(0)
    doc = Document(f)
    return "\n".join(p.text for p in doc.paragraphs)

def extract_text(f) -> str:
    ext = secure_filename(f.filename).rsplit(".", 1)[1].lower()

    if ext == "pdf":
        return extract_pdf(f)
    if ext == "docx":
        return extract_docx(f)

    # txt / rtf / html as plain text
    data = f.read()
    try:
        return data.decode("utf-8", errors="ignore")
    except AttributeError:
        # already str
        return data

# ===================================================== #
#                        ROUTES                         #
# ===================================================== #

@app.route("/")
def home():
    return render_template("index.html")

# ---------- MAIN PROCESS: REWRITE + SHOW --------------
@app.route("/rewrite", methods=["POST"])
def rewrite():

    resume_file = request.files.get("resume_file")
    job_text = request.form.get("job_description", "")

    if not resume_file:
        return "❌ No file uploaded", 400

    if not allowed_file(resume_file.filename):
        return "❌ Unsupported format", 400

    ext = resume_file.filename.rsplit(".",1)[1].lower()

    # ⭐ SAVE ORIGINAL for download
    resume_file.save("uploaded_original." + ext)

    resume_text = extract_text(resume_file)
    improved    = ai_rewrite_resume(resume_text, job_text)
    highlighted = highlight_keywords(improved, job_text)
    score       = compute_ats_score(improved, job_text)["score"]

    return render_template(
        "view_resume.html",
        rewritten_resume=highlighted,
        raw_resume=improved,
        job_text=job_text,
        score=score,
        original_ext=ext
    )

# ---------- IMPROVE AGAIN LOOP ------------------------
@app.route("/improve_again", methods=["POST"])
def improve_again():
    resume = request.form.get("resume_text", "")
    job = request.form.get("job_text", "")

    improved = ai_rewrite_resume(resume, job)
    highlighted = highlight_keywords(improved, job)
    score_data = compute_ats_score(improved, job)

    return jsonify({
        "rewritten": highlighted,
        "raw": improved,
        "score": score_data["score"]
    })

# ---------- CHATBOT API -------------------------------
@app.route("/chatbot_api", methods=["POST"])
def chatbot_api():
    msg = request.json.get("message", "")
    reply = genai.GenerativeModel("gemini-2.0-flash").generate_content(
        f"You are an ATS resume optimization assistant.\nUser says: {msg}"
    )
    return jsonify({"response": reply.text.strip()})

# ---------- DOWNLOAD AS SAME-TYPE FILE ----------------
@app.route("/download", methods=["POST"])
def download():
    text = request.form.get("resume_text", "")
    ext = request.form.get("ext", "docx").lower()

    # DOCX OUTPUT
    if ext == "docx":
        doc = Document()
        cleaned = text.replace("*", "").replace("**", "")
        for line in cleaned.split("\n"):
            doc.add_paragraph(line.strip())
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return send_file(
            buffer,
            as_attachment=True,
            download_name="ATS_Optimized_Resume.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    # PDF OUTPUT (simple text layout)
    if ext == "pdf":
        temp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
        c = canvas.Canvas(temp.name)

        y = 800
        for line in text.split("\n"):
            doc.add_paragraph(line)


        c.save()
        temp.seek(0)

        return send_file(
            temp.name,
            as_attachment=True,
            download_name="ATS_Optimized_Resume.pdf",
            mimetype="application/pdf"
        )

    # TXT / OTHER → return as plain text
    buffer = BytesIO()
    buffer.write(text.encode("utf-8", errors="ignore"))
    buffer.seek(0)
    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"ATS_Optimized_Resume.{ext}",
        mimetype="text/plain"
    )
@app.route("/download_txt", methods=["POST"])
def download_txt():
    text = request.form["resume_text"]
    buffer = BytesIO()
    buffer.write(text.encode("utf-8"))
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name="AI_Optimized_Resume.txt",
        mimetype="text/plain"
    )
@app.route("/download_doc", methods=["POST"])
def download_doc():

    text = request.form["resume_text"]

    # Create a .doc file manually (plain structured text)
    from io import BytesIO
    buffer = BytesIO()

    # Basic .doc header
    buffer.write(b'\xff\xfe')  # UTF16 BOM

    for line in text.split("\n"):
        buffer.write((line + "\r\n").encode("utf-16le"))

    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name="AI_Optimized_Resume.doc",
        mimetype="application/msword"
    )

# ===================================================== #
if __name__ == "__main__":
    app.run(debug=True)
